from __future__ import annotations

import logging
from pathlib import Path
import time
from typing import Any

import fitz
import pandas as pd
from docx import Document

from app.settings import Settings
from app.services.terminal_logging import log_blob, log_event

LOGGER = logging.getLogger(__name__)
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".xlsx"}


def parse_document(path: Path) -> dict[str, Any]:
    started_at = time.perf_counter()
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {extension}")
    log_event(
        LOGGER,
        "parser.dispatch.start",
        path=str(path),
        extension=extension,
        size_bytes=path.stat().st_size if path.exists() else None,
    )
    if extension == ".pdf":
        result = parse_pdf(path)
    elif extension == ".docx":
        result = parse_docx(path)
    elif extension == ".csv":
        result = parse_csv(path)
    elif extension == ".xlsx":
        result = parse_xlsx(path)
    else:
        result = parse_txt(path)
    log_blob(
        LOGGER,
        "parser.dispatch.result",
        result,
        path=str(path),
        elapsed_seconds=round(time.perf_counter() - started_at, 3),
    )
    return result


def parse_pdf(path: Path) -> dict[str, Any]:
    started_at = time.perf_counter()
    log_event(LOGGER, "parser.pdf.start", path=str(path))
    pages: list[dict[str, Any]] = []
    with fitz.open(path) as document:
        log_event(
            LOGGER,
            "parser.pdf.opened",
            path=str(path),
            page_count=document.page_count,
            metadata=document.metadata,
        )
        for index, page in enumerate(document, start=1):
            text = page.get_text("text").strip()
            pages.append({"page": index, "text": text})
            log_blob(
                LOGGER,
                "parser.pdf.page_text",
                text,
                path=str(path),
                page=index,
                character_count=len(text),
            )
    text = "\n\n".join(page["text"] for page in pages)
    metadata = {
        "parser": "pymupdf",
        "ocr_used": False,
        "ocr_engine": "",
        "ocr_confidence": None,
        "extracted_tables_count": 0,
        "extracted_images_count": count_pdf_images(path),
        "extraction_warnings": [],
    }
    if (
        Settings.ocr_enabled()
        and len(text.strip()) < Settings.ocr_min_text_characters()
    ):
        log_event(
            LOGGER,
            "parser.pdf.ocr_required",
            path=str(path),
            extracted_characters=len(text.strip()),
            minimum_characters=Settings.ocr_min_text_characters(),
        )
        ocr_result = parse_with_rapidocr(path)
        if ocr_result:
            log_blob(
                LOGGER,
                "parser.pdf.ocr_result",
                ocr_result,
                path=str(path),
                elapsed_seconds=round(time.perf_counter() - started_at, 3),
            )
            return ocr_result
        metadata["extraction_warnings"].append(
            "OCR was needed but the configured OCR pipeline was unavailable."
        )
    result = {
        "pages": pages,
        "text": text,
        "page_count": len(pages),
        "metadata": metadata,
    }
    log_blob(
        LOGGER,
        "parser.pdf.finish",
        result,
        path=str(path),
        elapsed_seconds=round(time.perf_counter() - started_at, 3),
    )
    return result


def parse_docx(path: Path) -> dict[str, Any]:
    started_at = time.perf_counter()
    log_event(LOGGER, "parser.docx.start", path=str(path))
    document = Document(path)
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    text = "\n".join(paragraphs)
    result = {
        "pages": [{"page": 1, "text": text}],
        "text": text,
        "page_count": 1,
        "metadata": parser_metadata("python-docx"),
    }
    log_blob(
        LOGGER,
        "parser.docx.finish",
        result,
        path=str(path),
        paragraph_count=len(paragraphs),
        table_count=len(document.tables),
        elapsed_seconds=round(time.perf_counter() - started_at, 3),
    )
    return result


def parse_txt(path: Path) -> dict[str, Any]:
    started_at = time.perf_counter()
    log_event(LOGGER, "parser.txt.start", path=str(path))
    text = path.read_text(encoding="utf-8", errors="ignore")
    result = {
        "pages": [{"page": 1, "text": text}],
        "text": text,
        "page_count": 1,
        "metadata": parser_metadata("plain-text"),
    }
    log_blob(
        LOGGER,
        "parser.txt.finish",
        result,
        path=str(path),
        elapsed_seconds=round(time.perf_counter() - started_at, 3),
    )
    return result


def parse_csv(path: Path) -> dict[str, Any]:
    started_at = time.perf_counter()
    log_event(LOGGER, "parser.csv.start", path=str(path))
    dataframe = pd.read_csv(path)
    text = dataframe.to_csv(index=False)
    result = {
        "pages": [{"page": 1, "text": text}],
        "text": text,
        "page_count": 1,
        "metadata": parser_metadata("pandas-csv", tables=1),
    }
    log_blob(
        LOGGER,
        "parser.csv.finish",
        result,
        path=str(path),
        rows=len(dataframe.index),
        columns=list(dataframe.columns),
        elapsed_seconds=round(time.perf_counter() - started_at, 3),
    )
    return result


def parse_xlsx(path: Path) -> dict[str, Any]:
    started_at = time.perf_counter()
    log_event(LOGGER, "parser.xlsx.start", path=str(path))
    frames = pd.read_excel(path, sheet_name=None)
    parts: list[str] = []
    for sheet_name, dataframe in frames.items():
        parts.append(f"Sheet: {sheet_name}")
        parts.append(dataframe.to_csv(index=False))
        log_blob(
            LOGGER,
            "parser.xlsx.sheet",
            dataframe.to_csv(index=False),
            path=str(path),
            sheet_name=sheet_name,
            rows=len(dataframe.index),
            columns=list(dataframe.columns),
        )
    text = "\n".join(parts)
    result = {
        "pages": [{"page": 1, "text": text}],
        "text": text,
        "page_count": 1,
        "metadata": parser_metadata("pandas-xlsx", tables=len(frames)),
    }
    log_blob(
        LOGGER,
        "parser.xlsx.finish",
        result,
        path=str(path),
        sheet_count=len(frames),
        elapsed_seconds=round(time.perf_counter() - started_at, 3),
    )
    return result


def parser_metadata(
    parser: str,
    *,
    tables: int = 0,
    images: int = 0,
    warnings: list[str] | None = None,
) -> dict[str, Any]:
    return {
        "parser": parser,
        "ocr_used": False,
        "ocr_engine": "",
        "ocr_confidence": None,
        "extracted_tables_count": tables,
        "extracted_images_count": images,
        "extraction_warnings": warnings or [],
    }


def count_pdf_images(path: Path) -> int:
    try:
        count = 0
        with fitz.open(path) as document:
            for page in document:
                count += len(page.get_images(full=True))
        log_event(LOGGER, "parser.pdf.image_count", path=str(path), image_count=count)
        return count
    except Exception:
        LOGGER.exception("PDF image counting failed. path=%s", path)
        return 0


def parse_with_rapidocr(path: Path) -> dict[str, Any] | None:
    started_at = time.perf_counter()
    log_event(LOGGER, "parser.ocr.start", path=str(path), engine=Settings.ocr_engine())
    try:
        import numpy as np
        from rapidocr import RapidOCR
    except Exception as exc:
        LOGGER.info("RapidOCR is unavailable. error=%s", exc)
        log_event(
            LOGGER,
            "parser.ocr.import_unavailable",
            path=str(path),
            error_type=type(exc).__name__,
            error=str(exc),
        )
        return None

    try:
        engine = RapidOCR()
        pages: list[dict[str, Any]] = []
        scores: list[float] = []
        with fitz.open(path) as document:
            for page_number, page in enumerate(document, start=1):
                pixmap = page.get_pixmap(dpi=200, alpha=False)
                image = np.frombuffer(pixmap.samples, dtype=np.uint8).reshape(
                    pixmap.height, pixmap.width, pixmap.n
                )
                output = engine(image)
                page_text = "\n".join(getattr(output, "txts", ()) or ())
                pages.append({"page": page_number, "text": page_text})
                scores.extend(float(score) for score in (output.scores or ()))
        text = "\n\n".join(page["text"] for page in pages)
        confidence = sum(scores) / len(scores) if scores else 0.0
        result = {
            "pages": pages,
            "text": text,
            "page_count": max(1, len(pages)),
            "metadata": {
                "parser": "rapidocr",
                "ocr_used": True,
                "ocr_engine": Settings.ocr_engine(),
                "ocr_confidence": round(confidence, 4),
                "extracted_tables_count": 0,
                "extracted_images_count": count_pdf_images(path),
                "extraction_warnings": (
                    [] if text.strip() else ["OCR returned no text."]
                ),
            },
        }
        log_blob(
            LOGGER,
            "parser.ocr.finish",
            result,
            path=str(path),
            recognised_lines=len(scores),
            elapsed_seconds=round(time.perf_counter() - started_at, 3),
        )
        return result
    except Exception as exc:
        LOGGER.exception("RapidOCR failed for %s", path)
        result = {
            "pages": [{"page": 1, "text": ""}],
            "text": "",
            "page_count": 1,
            "metadata": {
                "parser": "rapidocr",
                "ocr_used": True,
                "ocr_engine": Settings.ocr_engine(),
                "ocr_confidence": 0.0,
                "extracted_tables_count": 0,
                "extracted_images_count": count_pdf_images(path),
                "extraction_warnings": [f"OCR failed: {type(exc).__name__}"],
            },
        }
        log_blob(
            LOGGER,
            "parser.ocr.failed_result",
            result,
            level=logging.ERROR,
            path=str(path),
            error_type=type(exc).__name__,
            error=str(exc),
            elapsed_seconds=round(time.perf_counter() - started_at, 3),
        )
        return result


def ocr_health() -> dict[str, Any]:
    if not Settings.ocr_enabled():
        result = {
            "enabled": False,
            "available": False,
            "engine": Settings.ocr_engine(),
            "message": "OCR is disabled.",
        }
        log_blob(LOGGER, "parser.ocr_health", result)
        return result
    try:
        import rapidocr  # noqa: F401

        result = {
            "enabled": True,
            "available": True,
            "engine": Settings.ocr_engine(),
            "message": "RapidOCR import is available.",
        }
        log_blob(LOGGER, "parser.ocr_health", result)
        return result
    except Exception as exc:
        result = {
            "enabled": True,
            "available": False,
            "engine": Settings.ocr_engine(),
            "message": f"RapidOCR import failed: {type(exc).__name__}",
        }
        log_blob(LOGGER, "parser.ocr_health", result, level=logging.WARNING)
        return result
